"""Unit tests for document parsers."""

from __future__ import annotations

import pytest

from app.platform.providers.implementations.document_parser_factory import (
    CompositeDocumentParserProvider,
)
from app.platform.providers.implementations.docx_parser import DocxParserProvider
from app.platform.providers.implementations.plain_text_parser import PlainTextParserProvider
from app.platform.providers.implementations.pymupdf_parser import PyMuPDFParserProvider

pytestmark = pytest.mark.unit


def test_plain_text_parser_extracts_utf8() -> None:
    result = PlainTextParserProvider().parse(
        data=b"Hello parser",
        filename="hello.txt",
        content_type="text/plain",
    )
    assert result.text == "Hello parser"
    assert result.page_count == 1
    assert result.parser_name == "plain_text"
    assert result.parsed_document_version
    assert result.elements
    assert result.elements[0].element_type.value == "paragraph"


def test_markdown_parser_emits_headings_and_paragraphs() -> None:
    markdown = b"# Title\n\nBody paragraph.\n\n- item one\n- item two"
    result = PlainTextParserProvider().parse(
        data=markdown,
        filename="notes.md",
        content_type="text/markdown",
    )
    assert result.source_format.value == "markdown"
    element_types = [element.element_type.value for element in result.elements]
    assert "heading" in element_types
    assert "list" in element_types


def test_composite_rejects_unknown_type() -> None:
    parser = CompositeDocumentParserProvider()
    with pytest.raises(Exception) as exc_info:
        parser.parse(data=b"\x00\x01", filename="file.bin", content_type=None)
    assert "Unsupported" in str(exc_info.value)


def test_pymupdf_parser_reads_minimal_pdf() -> None:
    import fitz

    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "PDF text")
    pdf_bytes = document.tobytes()
    document.close()

    result = PyMuPDFParserProvider().parse(
        data=pdf_bytes,
        filename="sample.pdf",
        content_type="application/pdf",
    )
    assert result.page_count == 1
    assert "PDF text" in result.text
    assert result.parser_name == "pymupdf"
    assert result.elements
    assert result.parsed_document_version


def test_docx_parser_extracts_paragraphs() -> None:
    from io import BytesIO

    from docx import Document

    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Hello docx")
    document.add_paragraph("Second paragraph")
    document.save(buffer)

    result = DocxParserProvider().parse(
        data=buffer.getvalue(),
        filename="sample.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert result.page_count == 1
    assert "Hello docx" in result.text
    assert "Second paragraph" in result.text
    assert result.parser_name == "python_docx"
    assert result.elements


def test_composite_routes_docx() -> None:
    from io import BytesIO

    from docx import Document

    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Composite docx route")
    document.save(buffer)

    result = CompositeDocumentParserProvider().parse(
        data=buffer.getvalue(),
        filename="notes.docx",
        content_type=None,
    )
    assert result.parser_name == "python_docx"
    assert "Composite docx route" in result.text
