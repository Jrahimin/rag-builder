from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

OUT = "Dedicated_Hosted_Knowledge_Service_Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
RED = "9B1C1C"


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for idx, grid_col in enumerate(grid.gridCol_lst):
        grid_col.set(qn("w:w"), str(widths[idx]))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_keep(paragraph, with_next=False):
    p_pr = paragraph._p.get_or_add_pPr()
    element = OxmlElement("w:keepNext" if with_next else "w:keepLines")
    p_pr.append(element)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r = p2.add_run(text)
    set_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_section_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_keep(p, with_next=True)
    return p


def add_feature_block(doc, number, title, current, build, outcome):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{number}. {title}")
    set_font(r, size=12, color=DARK_BLUE, bold=True)
    rows = [
        ("Current foundation", current),
        ("Complete in development", build),
        ("Customer outcome", outcome),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [2300, 7060])
    for label, detail in rows:
        cells = table.add_row().cells
        shade(cells[0], LIGHT_BLUE)
        p1 = cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        set_font(r1, size=10, color=DARK_BLUE, bold=True)
        p2 = cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.12
        r2 = p2.add_run(detail)
        set_font(r2, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Styles: standard_business_brief preset.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    s = styles[name]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = True
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Number"):
    s = styles[name]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(11)
    s.paragraph_format.left_indent = Inches(0.50)
    s.paragraph_format.first_line_indent = Inches(-0.25)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.167

# Footer.
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fp.paragraph_format.space_before = Pt(0)
fp.paragraph_format.space_after = Pt(0)
r = fp.add_run("AI Platform Engine | Dedicated Hosted Knowledge Service Guide | Page ")
set_font(r, size=8.5, color=MUTED)
add_page_field(fp)

# Header.
hp = section.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = hp.add_run("PRODUCT DIRECTION")
set_font(r, size=8.5, color=MUTED, bold=True)

# Opening masthead.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("PRODUCT DEVELOPMENT GUIDE")
set_font(r, size=10, color=BLUE, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Dedicated Hosted Knowledge Service")
set_font(r, size=26, color=INK, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
r = p.add_run("The focused development path from today’s repository to a sellable private RAG service")
set_font(r, size=13, color=MUTED)

meta = doc.add_table(rows=3, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(meta, [1800, 7560])
meta_rows = [
    ("Primary offer", "Vendor-managed, dedicated document intelligence and grounded retrieval backend for B2B software products."),
    ("First customers", "Compliance, audit, and case-management SaaS vendors with document-heavy workflows."),
    ("Future option", "Supported self-hosted enterprise deployment after the hosted service is repeatable."),
]
for idx, (label, value) in enumerate(meta_rows):
    shade(meta.cell(idx, 0), LIGHT_BLUE)
    p1 = meta.cell(idx, 0).paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run(label)
    set_font(r1, size=10, color=DARK_BLUE, bold=True)
    p2 = meta.cell(idx, 1).paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(value)
    set_font(r2, size=10, color=INK)

doc.add_paragraph().paragraph_format.space_after = Pt(3)
add_callout(
    doc,
    "The decision",
    "Build one dependable private Knowledge API that an existing B2B product can call. The customer keeps its own application, user interface, customer relationships, and model-account ownership; APE operates the document ingestion, retrieval, evidence, and answer-generation backend.",
)

add_section_heading(doc, "1. What We Are Building")
p = doc.add_paragraph()
p.add_run("The product is not a generic AI platform, chatbot, vector database, or document OCR API. ").bold = True
p.add_run("It is a dedicated hosted service that gives a software company a reliable backend for turning its customers’ documents into searchable, grounded knowledge.")

add_section_heading(doc, "Core promise", level=2)
for text in [
    "Ingest the customer’s documents through a stable API.",
    "Process, index, and retrieve relevant evidence safely and asynchronously.",
    "Return grounded answers with page-level evidence the host product can display.",
    "Run in an isolated customer deployment, with the service operator responsible for the platform lifecycle.",
]:
    add_bullet(doc, text)

add_section_heading(doc, "Product boundary", level=2)
p = doc.add_paragraph("The host application remains responsible for its own UI, end-user login, workflow logic, and product-specific permissions. The knowledge service owns document processing, retrieval, evidence, and supported operations. For the first release, one project should represent one hard customer-data or authorization boundary.")

add_section_heading(doc, "2. The Dedicated Hosted Service Model")
p = doc.add_paragraph("The first commercial model is one isolated deployment per customer. This avoids the unnecessary development burden of a public multi-tenant SaaS control plane while giving customers a credible data-control story.")

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(table, [2100, 3600, 3660])
headers = ["Layer", "Service operator owns", "Customer owns"]
for i, h in enumerate(headers):
    shade(table.cell(0, i), LIGHT_BLUE)
    p = table.cell(0, i).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(h)
    set_font(r, size=10, color=DARK_BLUE, bold=True)
rows = [
    ("Platform", "API, workers, PostgreSQL/pgvector schema, Redis, object storage, monitoring, upgrades, backups", "Application integration and customer-facing workflow"),
    ("AI providers", "Supported provider integration and failure handling", "LLM and embedding credentials; usage cost"),
    ("Data", "Secure processing, indexing, delete/reprocess operations", "Data ownership, legal authority, retention policy"),
    ("Security boundary", "Service secrets, infrastructure security, operational response", "End-user identity, application permissions, project mapping"),
]
for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        p = cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.10
        r = p.add_run(value)
        set_font(r, size=9.5, color=INK, bold=(i == 0))

add_section_heading(doc, "3. The Minimum Development Work That Must Be Complete")
p = doc.add_paragraph("The repository already has the right foundation: project-scoped APIs, uploads, storage, parsing, chunking, embeddings, PostgreSQL/pgvector, hybrid retrieval, conversations, and worker processes. The work below is the minimum product-development gap to close before proposing the hosted service as a dependable paid offering.")

add_feature_block(
    doc, "1", "Durable document-processing jobs",
    "The service has queued document stages and Taskiq workers, but document status is also used as execution state. A worker interruption can leave a document stranded, and transient provider failures become terminal failures.",
    "Add a persisted processing-job record with stage, status, attempt count, idempotency key, lease/heartbeat, error type, and progress. Add automatic recovery of expired jobs and a deliberate retry policy for transient failures.",
    "Uploads become trustworthy: a customer can see whether a document is queued, running, ready, retrying, or permanently failed - and the service can recover without manual database intervention.",
)
add_feature_block(
    doc, "2", "First-class asynchronous job API",
    "Upload currently creates a document lifecycle, but not a stable product-level job resource for the integrating application.",
    "Return a job ID from upload/reprocess/reindex actions. Expose job status, progress, error reason, and safe retry. Support an idempotency header and a signed completion/failure webhook.",
    "The host product can update its own UI reliably without guessing from document states or polling indefinitely.",
)
add_feature_block(
    doc, "3", "Grounded answers with claim-level citations",
    "The current chat flow retrieves chunks and stores citation snapshots. That proves which context was supplied, not which source supports each answer claim.",
    "Require answer citations that map each factual claim or sentence to a stable document, page, and chunk. Add an evidence threshold and an explicit insufficient-evidence response when retrieval cannot support an answer.",
    "The customer can show defensible evidence in its product instead of presenting a generic AI answer with a source list.",
)
add_feature_block(
    doc, "4", "One production retrieval configuration",
    "The code has several provider adapters and development defaults such as hash embeddings and echo chat. The existing lexical reranker is not a learned reranker.",
    "Choose one supported hosted embedding/LLM route and one supported private/local route. Replace the lexical reranker with one measured learned reranker. Store a retrieval configuration version with each index run.",
    "Customers receive predictable quality and support, while the team avoids an unbounded model and provider matrix.",
)
add_feature_block(
    doc, "5", "Retrieval behavior for customer data",
    "Hybrid retrieval, metadata filters, and pgvector HNSW are present. Filtered approximate search and reranking behavior have not yet become a product-quality subsystem.",
    "Tune hybrid retrieval, candidate depth, reranking, metadata filters, context sizing, and filtered pgvector query behavior. Make re-embedding and reindexing controlled background operations tied to a configuration version.",
    "The service can handle real customer corpora, versions, and filters rather than only a technical demonstration corpus.",
)
add_feature_block(
    doc, "6", "Safe ingestion boundary",
    "The service limits file size and has parsers, but broad production file safety and acceptance handling are incomplete.",
    "Add file-signature and MIME validation, explicit supported-format rules, malware scanning, corrupt/password-protected file handling, and permanent failure codes. Maintain document purge/delete behavior as a job-backed lifecycle action.",
    "Customers know exactly what the service accepts and do not treat a failed parser as an unexplained AI problem.",
)
add_feature_block(
    doc, "7", "Dedicated-deployment configuration",
    "Most configuration is deployment-wide, which is acceptable for a dedicated instance, but hosted deployments must not silently use development substitutes.",
    "Require real provider credentials, validate provider connectivity/capabilities at startup, pin embedding dimensions, use deployment-level secrets, and make production provider selection explicit.",
    "Each customer deployment has a known, supportable configuration with no accidental hash embeddings or echo model behavior.",
)
add_feature_block(
    doc, "8", "Customer-controlled corpus operations",
    "Delete, reprocess, embed, and index concepts exist, but they should become safe asynchronous product operations.",
    "Provide job-backed reprocess, re-embed, reindex, delete, and purge actions with version-aware behavior. Prevent duplicate processing and preserve a clear history of what version is active.",
    "A customer can correct a bad extraction, change a supported model configuration, or remove data without operator-only intervention.",
)
add_feature_block(
    doc, "9", "Authorization hook only where needed",
    "Organization and project scoping already provide a useful base. Full end-user authorization inside a shared project does not yet exist.",
    "For version one, enforce one project per customer-data boundary. Add a signed entitlement/metadata-filter hook only if the first paying customer needs multiple permission levels within one corpus.",
    "The product has a clear, honest isolation model without prematurely building enterprise RBAC, SSO, or SCIM.",
)

add_section_heading(doc, "4. What Is Already Enough for the First Offer")
for text in [
    "PostgreSQL and pgvector as the single relational, vector, and keyword-search system.",
    "Organization API keys and project-scoped data access.",
    "Document upload, object storage, PDF/DOCX/TXT/Markdown processing, chunking, embedding, indexing, retrieval, and chat foundations.",
    "REST APIs and SSE streaming for the host application.",
    "One dedicated deployment per customer, operated by the service provider.",
    "Customer-owned model credentials and usage costs.",
]:
    add_bullet(doc, text)

add_callout(
    doc,
    "Initial content scope",
    "Contractually support digital PDFs, DOCX, TXT, and Markdown first. Treat complex tables, handwriting, scanned documents, and Bangla OCR as special requirements until a representative corpus and a certified parsing/OCR path exist.",
)

add_section_heading(doc, "5. Deliberately Not Required Before Selling")
p = doc.add_paragraph("These items may be useful later, but building them now delays the product without improving the first dedicated customer’s outcome.")
for text in [
    "Public multi-tenant SaaS control plane, self-service sign-up, free tier, or consumption billing.",
    "Customer dashboard beyond a small operator interface for jobs and corpus status.",
    "Python or TypeScript SDKs; REST and OpenAPI are enough initially.",
    "MCP, agents, multi-agent workflows, GraphRAG, fine-tuning, voice, or a no-code builder.",
    "Kubernetes, multi-region infrastructure, or active-active deployment.",
    "Multiple vector databases or a provider marketplace.",
    "Large connector catalog; the host application can upload documents initially.",
    "SSO, SCIM, enterprise RBAC, and complex per-user ACLs unless demanded by the first contract.",
]:
    add_bullet(doc, text)

add_section_heading(doc, "6. Development Sequence")
p = doc.add_paragraph("The sequence matters. Do not improve peripheral features before the processing and evidence path is dependable.")
for text in [
    "Make document processing durable: job record, retries, leases, recovery, and idempotency.",
    "Expose the product-level job API and signed webhook callbacks.",
    "Lock one real model/embedding configuration and introduce a learned reranker.",
    "Implement claim-level citations and insufficient-evidence behavior.",
    "Make reprocess, reindex, re-embed, delete, and purge safe asynchronous operations.",
    "Finish file validation and supported-file failure handling.",
    "Add the entitlement/filter hook only if the first customer needs it.",
]:
    add_number(doc, text)

add_section_heading(doc, "7. Future Option: Supported Self-Hosted Service")
p = doc.add_paragraph("Self-hosting should be the second delivery model, not a parallel initial product. It becomes realistic after the dedicated hosted service is repeatable across several customers.")

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(table, [2800, 6560])
for i, h in enumerate(["Build in the hosted product first", "Then reuse it for self-hosted"]):
    shade(table.cell(0, i), LIGHT_BLUE)
    p = table.cell(0, i).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(h)
    set_font(r, size=10, color=DARK_BLUE, bold=True)
self_host_rows = [
    ("Immutable service images and pinned dependency versions", "Customer receives supported release bundles rather than source-code instructions."),
    ("Safe migrations, reindexing, rollback, and backup/restore", "Customer can upgrade without an operator manually repairing data."),
    ("Configuration validation and provider compatibility checks", "Customer can use an approved model/storage setup with predictable support."),
    ("Operational diagnostics and structured job history", "Vendor support can diagnose problems inside a customer-controlled environment."),
]
for left, right in self_host_rows:
    cells = table.add_row().cells
    for i, value in enumerate((left, right)):
        p = cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.10
        r = p.add_run(value)
        set_font(r, size=9.5, color=INK, bold=(i == 0))

add_callout(
    doc,
    "Decision rule for self-hosting",
    "Offer it only when a qualified buyer cannot use a dedicated managed instance because of cloud-account ownership, data residency, procurement, or air-gap requirements - and only after the hosted deployment process is repeatable.",
)

add_section_heading(doc, "8. Proposal-Ready Definition")
p = doc.add_paragraph("The service is ready to propose when the following product statement is true:")
add_callout(
    doc,
    "Ready-to-propose statement",
    "We can deploy a dedicated instance for your product, accept the agreed document types, process them reliably in the background, let your application search or ask grounded questions, return evidence linked to document pages, and operate the platform lifecycle while you retain ownership of data and model-provider costs.",
)

add_section_heading(doc, "Positioning Statement")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("For compliance and case-management software vendors, the Private Knowledge API is a dedicated document intelligence and grounded retrieval backend that helps them add reliable, citable AI search and answers to their existing products without building a RAG platform themselves or sending customer data through a shared AI data plane.")
set_font(r, size=12, color=INK, bold=True)

add_section_heading(doc, "Reference: Existing Repository Foundation", level=2)
p = doc.add_paragraph("This guide is based on the repository review. The platform already contains the main architectural pieces for the proposed service: scoped API routing, project-scoped repositories, asynchronous document processing, pgvector-backed retrieval, Docker development topology, and documented production gaps.")
src = doc.add_paragraph()
src.paragraph_format.space_after = Pt(0)
src.add_run("Repository sources: ")
add_hyperlink(src, "project context", "E:/python-projects/rag-builder/.cursor/rules/project-context.mdc")
src.add_run("; ")
add_hyperlink(src, "deployment architecture", "E:/python-projects/rag-builder/docs/architecture/deployment-architecture.md")
src.add_run("; ")
add_hyperlink(src, "background processing", "E:/python-projects/rag-builder/docs/architecture/background-processing.md")

doc.core_properties.title = "Dedicated Hosted Knowledge Service Guide"
doc.core_properties.subject = "Product direction and minimum development scope"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT)
