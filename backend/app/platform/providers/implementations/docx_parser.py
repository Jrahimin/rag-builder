"""Microsoft Word (.docx) text extraction via python-docx."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from app.platform.domain.language_detection import detect_language
from app.platform.providers.contracts.document_parser import (
    BaseDocumentParserProvider,
    ParsedDocument,
    ParsedElement,
    ParsedElementType,
    SourceFormat,
)
from app.platform.providers.errors import ProviderError
from app.platform.providers.implementations.parsed_element_builder import finalize_elements

_PARSER_NAME = "python_docx"
_PARSER_VERSION = "1.1.2"
_DOCX_EXTENSIONS = {".docx"}
_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class DocxParserProvider(BaseDocumentParserProvider):
    """Extract structured paragraph, heading, list, and table elements from DOCX."""

    def parse(
        self,
        *,
        data: bytes,
        filename: str,
        content_type: str | None,
        ocr_lang: str | None = None,
    ) -> ParsedDocument:
        del filename, content_type, ocr_lang
        if not data:
            return ParsedDocument(
                text="",
                page_count=0,
                parser_name=_PARSER_NAME,
                parser_version=_PARSER_VERSION,
                source_format=SourceFormat.DOCX,
                parser_confidence=0.0,
            )

        try:
            document = DocxDocument(BytesIO(data))
            elements = _extract_elements(document)
            text, finalized = finalize_elements(elements)
        except PackageNotFoundError as exc:
            msg = "Input is not a valid DOCX document."
            raise ProviderError(msg, provider_name=_PARSER_NAME) from exc
        except Exception as exc:
            msg = "Failed to parse DOCX document."
            raise ProviderError(msg, provider_name=_PARSER_NAME) from exc

        heading_count = sum(
            1 for element in finalized if element.element_type == ParsedElementType.HEADING
        )
        confidence = 0.9 if heading_count else 0.75 if finalized else 0.0
        language_result = detect_language(text)

        return ParsedDocument(
            text=text,
            page_count=1 if text.strip() else 0,
            parser_name=_PARSER_NAME,
            parser_version=_PARSER_VERSION,
            elements=finalized,
            source_format=SourceFormat.DOCX,
            parser_confidence=confidence,
            language=language_result.primary_language,
            structure_hints={
                "heading_count": heading_count,
                "language_confidence": language_result.confidence,
                "languages": language_result.languages,
                "is_mixed": language_result.is_mixed,
            },
        )

    @staticmethod
    def supports(filename: str, content_type: str | None) -> bool:
        lower = filename.lower()
        if any(lower.endswith(ext) for ext in _DOCX_EXTENSIONS):
            return True
        return content_type == _DOCX_CONTENT_TYPE


def _extract_elements(document: Any) -> list[ParsedElement]:
    """Collect structured elements in document order."""
    elements: list[ParsedElement] = []
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            element = _paragraph_to_element(paragraph)
            if element is not None:
                elements.append(element)
        elif tag == qn("w:tbl"):
            table_text = _table_to_text(document, child)
            if table_text:
                elements.append(
                    ParsedElement(
                        text=table_text,
                        element_type=ParsedElementType.TABLE,
                        page_start=1,
                        page_end=1,
                    )
                )
    return elements


def _paragraph_to_element(paragraph: Paragraph) -> ParsedElement | None:
    text = paragraph.text.strip()
    if not text:
        return None
    style_name = (paragraph.style.name or "").lower() if paragraph.style is not None else ""
    if style_name.startswith("heading"):
        level = _heading_level_from_style(style_name)
        return ParsedElement(
            text=text,
            element_type=ParsedElementType.HEADING,
            page_start=1,
            page_end=1,
            heading_level=level,
        )
    if _is_list_paragraph(paragraph):
        return ParsedElement(
            text=text,
            element_type=ParsedElementType.LIST,
            page_start=1,
            page_end=1,
        )
    return ParsedElement(
        text=text,
        element_type=ParsedElementType.PARAGRAPH,
        page_start=1,
        page_end=1,
    )


def _heading_level_from_style(style_name: str) -> int:
    digits = "".join(char for char in style_name if char.isdigit())
    if digits:
        return int(digits)
    return 1


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name or "").lower() if paragraph.style is not None else ""
    if "list" in style_name or "bullet" in style_name:
        return True
    return text_has_list_prefix(paragraph.text)


def text_has_list_prefix(text: str) -> bool:
    stripped = text.lstrip()
    return stripped.startswith(("- ", "* ", "• ")) or (
        len(stripped) > 2 and stripped[0].isdigit() and stripped[1:3] in {". ", ") "}
    )


def _table_to_text(document: Any, table_element: Any) -> str:
    from docx.table import Table

    table = Table(table_element, document)
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)
